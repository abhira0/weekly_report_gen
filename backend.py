import os

import docx
import docx2pdf
import yaml
from docx.document import Document as DOCument
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches, Pt
from docx.table import _Cell, _Column, _Row
from termcolor import cprint

with open("config.yml", "r") as f:
    config_yml = yaml.load(f, Loader=yaml.FullLoader)


class YouTil:
    @staticmethod
    def cleanPathName(text):
        arr = ["\\", "/", "<", ">", "|", '"', "?", "*", ":"]
        for pattern in arr:
            text = text.replace(pattern, "")
        return text

    @staticmethod
    def makedir(path: str):
        if not os.path.exists(path):
            os.mkdir(path)


class Week:
    def __init__(self, week_no: int = None) -> None:
        self.week_no = week_no if week_no else int(input("Week Number        : "))

    def getInfo(self):
        self._setWeekDates()
        self.completed_activities = self._getActivities("Activities Completed")
        self.in_progress_activities = self._getActivities("In Progress")
        self.planned_activities = self._getActivities("Plan for Next Week")
        self.ultimatum = self._getUltimatum()
        print(self.ultimatum)
        self._dumpUltimatum()

    def _setWeekDates(self):
        print("-" * 15 + " Dates " + "-" * 15)
        self.starting_date = input("Week Starting Date : ")
        self.ending_date = input("Week Ending Date   : ")

    def _getActivities(self, activity_type):
        print("-" * 15 + f" {activity_type} " + "-" * 15)
        try:
            count = int(input("Number of tasks: "))
        except:
            count = 0
        return [self._getTaskDetails(c + 1) for c in range(count)]

    def _getTaskDetails(self, count):
        task_description = input(f"{count}. Task Description: ")
        task_date = input(f"{count}. Date            : ")
        return {"task_description": task_description, "task_date": task_date}

    def _getUltimatum(self):
        return {
            "week_no": self.week_no,
            "week_dates": {"start": self.starting_date, "end": self.ending_date},
            "completed_activities": self.completed_activities,
            "in_progress_activities": self.in_progress_activities,
            "planned_activities": self.planned_activities,
        }

    def _dumpUltimatum(self):
        cwd = os.path.dirname(os.path.realpath(__file__))
        path_ = f"{cwd}\\weekConfigs"
        YouTil.makedir(path_)
        with open(f"{path_}\\{self.week_no}.yml", "w") as f:
            f.write(yaml.dump(self.ultimatum, sort_keys=False))
        cprint(f"yml file saved at .\\weekConfigs\\{self.week_no}.yml", "green")


class TableEntries:
    def __init__(self, doc: DOCument, ultimatum) -> None:
        self.doc = doc
        self.ultimatum = ultimatum
        self.table = self.doc.tables[0]

    def writeProjectDetails(self):
        # Style - Project ID
        C_ProjectID = self.doc.styles.add_style("C_ProjectID", WD_STYLE_TYPE.PARAGRAPH)
        C_ProjectID.font.size = Pt(20)
        C_ProjectID.font.bold = True
        # Style -  Project Title
        C_ProjectID = self.doc.styles.add_style(
            "C_ProjectTitle", WD_STYLE_TYPE.PARAGRAPH
        )
        C_ProjectID.font.size = Pt(15)
        C_ProjectID.font.bold = True
        # Project ID
        para: _Cell = self.table.cell(0, 2).paragraphs[0]
        para.style = self.doc.styles["C_ProjectID"]
        para.text = config_yml["project"]["id"]
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # Project Title
        para: _Cell = self.table.cell(1, 2).paragraphs[0]
        para.style = self.doc.styles["C_ProjectTitle"]
        para.text = config_yml["project"]["title"]

    def writeWeekDates(self):
        # Style - Week dates
        C_ProjectID = self.doc.styles.add_style("C_WeekDates", WD_STYLE_TYPE.PARAGRAPH)
        C_ProjectID.font.size = Pt(12)
        C_ProjectID.font.italic = True
        # Week Starting Date
        para: _Cell = self.table.cell(2, 2).paragraphs[0]
        para.style = self.doc.styles["C_WeekDates"]
        para.text = self.ultimatum["week_dates"]["start"]
        # Week Ending Date
        para: _Cell = self.table.cell(2, 5).paragraphs[0]
        para.style = self.doc.styles["C_WeekDates"]
        para.text = self.ultimatum["week_dates"]["end"]

    def delete_paragraph(self, paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None

    def writeTeamDetails(self):
        # Style - Names, SRNs
        C_ProjectID = self.doc.styles.add_style("C_NamesSRNs", WD_STYLE_TYPE.PARAGRAPH)
        C_ProjectID.font.size = Pt(12)
        # C_ProjectID.font.bold = True
        # Student names
        para: _Cell = self.table.cell(3, 0).paragraphs[1]
        para.style = self.doc.styles["C_NamesSRNs"]
        para.text = "\n".join([i["name"] for i in config_yml["team"]])
        self.delete_paragraph(self.table.cell(3, 0).paragraphs[4])
        self.delete_paragraph(self.table.cell(3, 0).paragraphs[3])
        self.delete_paragraph(self.table.cell(3, 0).paragraphs[2])
        # Student SRNs
        para: _Cell = self.table.cell(3, 4).paragraphs[1]
        para.style = self.doc.styles["C_NamesSRNs"]
        para.text = "\n".join([i["srn"] for i in config_yml["team"]])
        self.delete_paragraph(self.table.cell(3, 4).paragraphs[4])
        self.delete_paragraph(self.table.cell(3, 4).paragraphs[3])
        self.delete_paragraph(self.table.cell(3, 4).paragraphs[2])

    def addRowAt(self, table, ind: int):
        table.add_row()
        row_ind = table.rows[ind - 1]  # for example
        new_row = table.rows[-1]
        row_ind._tr.addnext(new_row._tr)
        table.cell(ind, 1).merge(table.cell(ind, 2))
        table.cell(ind, 4).merge(table.cell(ind, 5))

    def writeActivities(self):
        # Style - Activities
        C_ProjectID = self.doc.styles.add_style("C_Activity", WD_STYLE_TYPE.PARAGRAPH)
        C_ProjectID.font.size = Pt(12)
        # Completed Activities
        row_no = 6
        activity_list = self.ultimatum["completed_activities"]
        [self.writeATask(c, row_no + c, ele) for c, ele in enumerate(activity_list)]
        # In Progress Activities
        row_no += 2 + len(activity_list)
        activity_list = self.ultimatum["in_progress_activities"]
        [self.writeATask(c, row_no + c, ele) for c, ele in enumerate(activity_list)]
        # Planned Activities
        row_no += 2 + len(activity_list)
        activity_list = self.ultimatum["planned_activities"]
        [self.writeATask(c, row_no + c, ele) for c, ele in enumerate(activity_list)]

    def writeATask(self, sl_no: int, row_no: int, task_details: dict):
        self.addRowAt(self.table, row_no)
        # Serial Number
        para: _Cell = self.table.cell(row_no, 0).paragraphs[0]
        para.style = self.doc.styles["C_Activity"]
        para.text = str(sl_no + 1)
        # Task Description
        para: _Cell = self.table.cell(row_no, 1).paragraphs[0]
        para.style = self.doc.styles["C_Activity"]
        para.text = task_details["task_description"]
        # Date
        para: _Cell = self.table.cell(row_no, 3).paragraphs[0]
        para.style = self.doc.styles["C_Activity"]
        para.text = task_details["task_date"]


class Converter:
    def __init__(self, week_no: int = None) -> None:
        self.cwd = os.path.dirname(os.path.realpath(__file__))
        self.week_no = week_no if week_no else int(input("Week No: "))
        self._setUltimatum()
        self.file_name = "week" + str(self.ultimatum["week_no"])

    def _setUltimatum(self):
        path_ = f"{self.cwd}\\weekConfigs"
        if os.path.exists(f"{path_}\\{self.week_no}.yml"):
            with open(f"{path_}\\{self.week_no}.yml", "r") as f:
                self.ultimatum = yaml.load(f, Loader=yaml.FullLoader)
        else:
            cprint(
                "Please build the week YAML (by using -b flag) file initially before creating the docx|pdf",
                "red",
            )
            quit(0)

    def createDocx(self, verbose=False):
        self.doc: DOCument = docx.Document(f"{self.cwd}\\template\\template.docx")
        self._createTable()
        YouTil.makedir(f"{self.cwd}\\reports\\")
        self.doc.save(f"{self.cwd}\\reports\\{self.file_name}.docx")
        if verbose:
            cprint(f"docx saved at .\\reports\\{self.file_name}.docx", "green")

    def flushDocx(self):
        os.remove(f"{self.cwd}\\reports\\{self.file_name}.docx")

    def _createTable(self):
        te = TableEntries(self.doc, self.ultimatum)
        te.writeProjectDetails()
        te.writeWeekDates()
        te.writeTeamDetails()
        te.writeActivities()

    def saveAsPDF(self):
        self.createDocx()
        path1 = f"{self.cwd}\\reports\\{self.file_name}.docx"
        path2 = f"{self.cwd}\\reports\\{self.file_name}.pdf"
        print("Converting DOCX to PDF")
        docx2pdf.convert(path1, path2)
        cprint(f"docx saved at .\\reports\\{self.file_name}.docx", "green")
